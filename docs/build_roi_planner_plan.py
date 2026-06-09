"""Generate the ROI Planner execution plan + Tuesday showcase doc."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0B, 0x2E, 0x4F)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x15, 0x80, 0x3D)
AMBER = RGBColor(0xB4, 0x6B, 0x00)
RED = RGBColor(0xB4, 0x22, 0x22)

doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_title(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_para(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_num(text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


# ---------- TITLE ----------
add_title("AI ROI Planner — Execution Plan & Tuesday Showcase")
add_subtitle("CoAction Labs AI Hub  |  Prepared by Ramana Narayanan  |  Tuesday Review")

# ---------- 1. ONE-LINE FRAMING ----------
add_h1("1. The One-Liner")
add_para(
    "I have a working AI ROI Planner prototype. The plan is to fold it into the CoAction Labs "
    "AI Hub as a permissioned /roi-planner module on our standard Node + Express + Postgres + "
    "Drizzle + Okta stack. On Tuesday I will demo the prototype, walk through the integration "
    "plan, and ask for approval to start Week 1."
)

# ---------- 2. WHERE WE ARE TODAY ----------
add_h1("2. Where We Are Today")

add_h2("Already built")
add_bullet("Portal shell — React + TS + Vite + Tailwind + shadcn/ui, EnterpriseHeader, AuthContext, SortableTable with themed scroll affordance, the full shadcn UI kit.")
add_bullet("Reference pages following the same pattern I will mirror — ProjectCatalogue, MyAgents, Analytics, PoliciesGovernance.")
add_bullet("ROI Planner prototype — all 9 tabs (Dashboard, Executive, Run rate, Initiatives, Roadmap, Heatmap, HC harvest, Goals, Settings) with the full financial model, methodology panel, dark/light theme, and view-only mode.")
add_bullet("Tuesday brief document (ROI_Planner_Tuesday_Brief.docx).")

add_h2("Not built yet")
add_bullet("Backend — no apps/api/ project exists. No Express, no Drizzle, no Okta middleware.")
add_bullet("Database — no migrations applied to labs_ai_hub for catalog.initiatives or any ROI tables.")
add_bullet("Integration — the prototype is a standalone 1.6k-line file that uses window.storage. Not yet inside src/pages/ and not yet decomposed.")
add_bullet("Tests — zero coverage on the financial math today. This is the single biggest risk.")
add_bullet("Permissions — Viewer / Contributor / Admin model is designed but not enforced.")

# ---------- 3. THE PLAN ----------
add_h1("3. The Plan — 4 Weeks to Pilot")

add_h2("Week 1 — Foundation")
add_bullet("Scaffold apps/api/ (Node + Express + TypeScript).")
add_bullet("Wire Okta JWT middleware. Roles: Viewer, Contributor, Admin.")
add_bullet("Stand up Drizzle against labs_ai_hub.")
add_bullet("Write migrations for catalog.initiatives, catalog.initiative_phases, catalog.initiative_roi_models, catalog.initiative_goals, ops.reference_employee_types.")
add_bullet("Apply lifecycle.workflow_states rows for the initiative state machine (draft → planned → active → completed).")
add_bullet("Confirm reuse of existing integration.vendor_platforms, engagement.comments, audit.events, iam.")

add_h2("Week 2 — Module + Math")
add_bullet("Build apps/api/src/modules/roi/ — routes, controller, service (team scoping), repo, dto (zod).")
add_bullet("Port the prototype's pure functions into roi.calc.ts (shared with the client).")
add_bullet("Vitest suite on roi.calc with locked baseline numbers. This unblocks any exec demo.")
add_bullet("Drop the prototype into src/pages/RoiPlanner/ and replace window.storage with src/api/roi.ts (fetch wrapper that attaches the Okta access token).")
add_bullet("Add /roi-planner route in App.tsx and the nav entry in EnterpriseHeader.")

add_h2("Week 3 — Decompose + Polish + Pilot Start")
add_bullet("Break the 1.6k-line file into RoiPlannerPage.tsx + tabs/*.tsx + components/*.tsx + calc/ + types.ts.")
add_bullet("Apply the portal design system — shadcn primitives, Tailwind palette, EnterpriseHeader, PageHeader.")
add_bullet("Remove the unauthenticated share/import blob (or replace with signed share links).")
add_bullet("Wire view-only mode to the Okta role, not a manual toggle.")
add_bullet("Seed sample CoAction initiatives for the pilot team.")
add_bullet("Cross-test client calc/ against server roi.calc.ts — same inputs, identical outputs.")
add_bullet("Pilot with one team (Growth Protocol or AI Engineering).")

add_h2("Week 4+ — Org Rollout")
add_bullet("Open to the org behind the Okta group.")
add_bullet("Cross-link initiatives with Project Catalogue.")
add_bullet("Vendor auto-population from integration.vendor_platforms.")
add_bullet("Nightly refresh job for analytics.initiative_yearly_rollups.")
add_bullet("Audit trigger on catalog.initiatives → audit.events.")

# ---------- 4. WORKSTREAM TABLE ----------
add_h1("4. Workstream Summary")
add_table(
    headers=["Workstream", "Status", "Week"],
    rows=[
        ["Portal shell (header, auth, table, UI kit)", "Done", "—"],
        ["ROI Planner prototype (UI + financial model)", "Done", "—"],
        ["apps/api/ scaffold + Okta middleware", "Not started", "1"],
        ["Drizzle schema + migrations on labs_ai_hub", "Not started", "1"],
        ["roi module (routes, service, repo, dto)", "Not started", "2"],
        ["Vitest coverage on roi.calc", "Not started", "2"],
        ["Port prototype into src/pages/RoiPlanner/", "Not started", "2"],
        ["Replace window.storage with src/api/roi.ts", "Not started", "2"],
        ["Decompose 1.6k-line file", "Not started", "3"],
        ["Design-system polish", "Not started", "3"],
        ["Role-based view-only + team scoping", "Not started", "3"],
        ["Pilot with one team", "Not started", "3"],
        ["Project Catalogue cross-linking", "Not started", "4+"],
        ["Nightly rollup job + audit trigger", "Not started", "4+"],
    ],
    col_widths=[3.4, 1.4, 0.8],
)

# ---------- 5. WHAT I WILL SHOWCASE ON TUESDAY ----------
add_h1("5. What I Will Showcase on Tuesday")
add_subtitle("~15 min live demo, ~10 min plan, ~5 min Q&A")

add_h2("Live Demo Flow")
add_num("Land on Dashboard with sample CoAction initiatives loaded (Vega, CoSave, Cortex).")
add_num("Executive view — show the FTE / BPO / Roadmap split. Click the Build vs Subscribe tile to demonstrate live filtering.")
add_num("Drill-down into any initiative → open the calculation methodology panel. This is the trust-builder for finance.")
add_num("Run-rate view — toggle the year filter, walk the Completed / Active / Planned / Roadmap pipeline.")
add_num("Investment arc — 5-year savings vs costs, breakeven marker, cumulative net.")
add_num("HC harvest timeline — quarterly FTE positions, grouped by Domain, expand a business function.")
add_num("Roadmap Gantt — savings phases highlighted in green, 'now' line visible.")
add_num("Impact heatmap — switch sizing between Net ROI, Savings, and HC.")
add_num("Goals tab — planned vs realized progress with roadmap-opportunity overlay.")
add_num("Add an initiative live — fill the form, ROI preview updates instantly. Flip to view-only mode to demonstrate the permission model.")

add_h2("Integration Plan Walkthrough")
add_bullet("Stack: Node + Express + Postgres + Drizzle + Okta — same as the rest of the Hub.")
add_bullet("Schema mapping: catalog.initiatives (+ children) for assets, analytics.initiative_yearly_rollups for dashboards, integration.vendor_platforms reused, lifecycle.workflow_states for state machine, engagement.comments for discussion, audit.events for trail.")
add_bullet("Roles: Viewer / Contributor / Admin enforced in Okta middleware.")
add_bullet("Frontend: /roi-planner route, src/pages/RoiPlanner/, src/api/roi.ts wrapper.")
add_bullet("Rollout: Week 1 foundation, Week 2 module + math, Week 3 decompose + pilot, Week 4+ org-wide.")

add_h2("Risks I Will Name Out Loud")
add_bullet("Financial math has no tests today → Vitest before any exec demo.")
add_bullet("1.6k-line single file → decompose in Week 3.")
add_bullet("Roadmap 'estimated' savings could be misread as committed → distinct color + sub-labels in copy.")
add_bullet("Unauthenticated share/import blob → remove for production, replace with signed share links if needed.")
add_bullet("Local date math could drift across timezones → normalize to UTC in roi.calc, aggregate at month granularity.")

# ---------- 6. THE ASKS ----------
add_h1("6. What I Need From the Meeting")
add_num("Approval to scaffold apps/api/src/modules/roi/ on the agreed Node + Express + Drizzle + Okta stack.")
add_num("Sign-off on the schema mapping into catalog / analytics / integration / lifecycle / engagement.")
add_num("A pilot team commitment — Growth Protocol or AI Engineering preferred (both already in Project Catalogue).")
add_num("Agreement that financial math gets Vitest coverage before any exec sees the numbers.")
add_num("Naming sign-off — 'ROI Planner' vs 'AI Investment Tracker' vs alternative.")

# ---------- 7. SUCCESS CRITERIA ----------
add_h1("7. Definition of Done")
add_bullet("Pilot team is using /roi-planner against labs_ai_hub Postgres, authenticated via Okta.")
add_bullet("Every initiative has an owner, a status, and a non-null Net ROI.")
add_bullet("roi.calc has Vitest coverage on every public function, baselines locked.")
add_bullet("Executive view loads under 1.5 seconds with the pilot team's full dataset.")
add_bullet("Audit trail captures every initiative create / update / delete.")
add_bullet("Project Catalogue and ROI Planner share the same initiative records — single source of truth.")

# ---------- FOOTER ----------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
run = p.add_run("Prepared by Ramana Narayanan  |  CoAction Labs — AI Engineering  |  Companion to ROI_Planner_Tuesday_Brief.docx")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = GRAY

out = "/Users/navi/Work/smart-agent-foundry-main/docs/ROI_Planner_Plan_and_Showcase.docx"
doc.save(out)
print("WROTE:", out)
