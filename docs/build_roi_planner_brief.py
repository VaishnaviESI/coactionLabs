"""Generate the Tuesday meeting brief for the AI ROI Planner integration."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

NAVY = RGBColor(0x0B, 0x2E, 0x4F)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x15, 0x80, 0x3D)
RED = RGBColor(0xB4, 0x22, 0x22)

doc = Document()

# Base styles
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_subtitle(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.italic = True
    run.font.color.rgb = GRAY


def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY


def add_para(text, bold=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            r = p.add_run(str(val))
            r.font.size = Pt(10)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)


# ---------- TITLE ----------
add_title("AI ROI Planner — Integration Brief")
add_subtitle("CoAction Labs AI Hub  |  Tuesday Review  |  Prepared for stakeholders")

# ---------- 1. EXECUTIVE SUMMARY ----------
add_h1("1. Executive Summary")
add_para(
    "I have a working prototype of an AI ROI Planner — a single tool that captures every AI "
    "initiative across CoAction, models its 4-year financial impact, and presents it to "
    "executives, finance, and initiative owners with role-appropriate views."
)
add_para(
    "It complements the existing Project Catalogue: that tells us WHAT we are doing; the ROI "
    "Planner tells us WHAT IT IS WORTH. I am proposing to integrate it into the CoAction Labs "
    "AI Hub as a permissioned module, on the same Node + Express + Postgres + Okta stack we are "
    "standardizing on for the rest of the platform."
)

# ---------- 2. WHAT IS IN THE PROTOTYPE ----------
add_h1("2. What the Prototype Already Does")
add_bullet("Initiative CRUD with phases, vendors, employee types, HITL costs, and per-transaction savings.")
add_bullet("Executive view: tile-based summary by FTE/BPO/Roadmap with click-through drill-downs.")
add_bullet("Run-rate view: pipeline by status (Completed / Active / Planned / Roadmap) with year filter.")
add_bullet("Investment arc: 5-year savings vs costs with cumulative net and breakeven marker.")
add_bullet("HC harvest timeline: quarterly view of FTE positions harvested, grouped by Domain / LOB / AI Type, with business-function drill-down.")
add_bullet("Roadmap Gantt with savings-phase highlighting and 'now' line.")
add_bullet("Impact heatmap sized by Net ROI, Savings, or HC.")
add_bullet("Goals & benchmarking with planned vs realized progress and roadmap-opportunity overlay.")
add_bullet("Dark/Light themes, view-only mode, share/import via encoded blob.")
add_bullet("Calculation methodology panel surfaced in every drill-down — auditability built in.")

# ---------- 3. WHY IT BELONGS IN THE AI HUB ----------
add_h1("3. Why It Belongs in the AI Hub")
add_bullet("Answers the single question every executive asks about AI: 'what are we getting back?'")
add_bullet("Same audience and same auth boundary as the Project Catalogue, Academy, and Governance tiles.")
add_bullet("Reuses concepts already in the Hub: initiatives, owners, teams, LOBs, vendors.")
add_bullet("Surfaces governance signals (no savings phase, missing owner, stale ROI) that today are invisible.")

# ---------- 4. INTEGRATION PLAN ----------
add_h1("4. Integration Plan (Node + Express + Postgres + Okta)")

add_h2("4.1 Database")
add_para(
    "No new bounded contexts required. The component maps cleanly into the labs_ai_hub schema:"
)
add_table(
    headers=["Concept", "Schema → Table", "Notes"],
    rows=[
        ["Initiative", "catalog.initiatives", "New asset type alongside agents, prompts, tools."],
        ["Phases", "catalog.initiative_phases", "Child rows, cascade on delete."],
        ["ROI model + headcount rows", "catalog.initiative_roi_models", "Versioned for reproducibility."],
        ["HITL + per-tx model", "jsonb on catalog.initiatives", "Sparse, schema-flexible."],
        ["Roadmap opportunity", "jsonb on catalog.initiatives", "Only populated for 'roadmap' status."],
        ["Goals & targets", "catalog.initiative_goals", "Owned by the catalog domain."],
        ["Vendors / platforms", "integration.vendor_platforms", "Reused across the Hub."],
        ["Employee types & costs", "ops.reference_employee_types", "Admin-managed reference data."],
        ["Yearly rollups", "analytics.initiative_yearly_rollups", "Materialized view, refreshed nightly."],
        ["Status transitions", "lifecycle.workflow_states", "draft → planned → active → completed."],
        ["Comments", "engagement.comments", "Polymorphic on subject_type='initiative'."],
        ["Audit trail", "audit.events", "Auto-populated via trigger."],
    ],
    col_widths=[1.7, 2.4, 2.6],
)

add_h2("4.2 API Module (apps/api/src/modules/roi/)")
add_para("Mirrors the catalog/agents template:")
add_bullet("roi.routes.ts — Express routes")
add_bullet("roi.controller.ts — HTTP layer")
add_bullet("roi.service.ts — business logic + team scoping")
add_bullet("roi.repo.ts — Drizzle queries")
add_bullet("roi.calc.ts — ported pure financial math, shared with the client")
add_bullet("roi.schema.ts — Drizzle table definitions")
add_bullet("roi.dto.ts — zod validators")
add_bullet("__tests__/roi.calc.spec.ts — Vitest coverage of the math")

add_h2("4.3 Auth & Permissions (Okta JWT middleware)")
add_para("Replaces Postgres RLS with explicit middleware — easier to debug, same security guarantees.")
add_table(
    headers=["Role", "Capability"],
    rows=[
        ["Viewer", "Read all initiatives in their team. Dashboards only."],
        ["Contributor", "Create / edit initiatives they own. View team data."],
        ["Admin", "Full access within their team. Manages settings, vendors, employee types."],
    ],
    col_widths=[1.2, 5.5],
)

add_h2("4.4 Frontend Layout")
add_bullet("New route: /roi-planner behind the same auth as the rest of the Hub.")
add_bullet("Page lives under src/pages/RoiPlanner/, decomposed from the current single 1.6k-line file.")
add_bullet("Data access via src/api/roi.ts — fetch wrapper that attaches the Okta access token (no direct fetch in components).")
add_bullet("Reuses EnterpriseHeader, Tailwind palette, and shadcn primitives to match the rest of the portal.")

# ---------- 5. RISKS ----------
add_h1("5. Risks & Mitigations")
add_table(
    headers=["Risk", "Mitigation"],
    rows=[
        ["Financial math has no test coverage today.", "Add Vitest suite around roi.calc before any exec demo."],
        ["1.6k-line single file is hard to maintain.", "Decompose into modules in Week 2."],
        ["Roadmap 'estimated' savings could be misread as committed.", "Distinct color treatment + sub-labels already in the UI; reinforce in copy."],
        ["Prototype's share/import blob is unauthenticated.", "Remove for production; replace with signed share links if needed."],
        ["Local date math could drift across timezones.", "Aggregate at month granularity (already the case); add UTC normalization in roi.calc."],
    ],
    col_widths=[3.3, 3.4],
)

# ---------- 6. ROLLOUT TIMELINE ----------
add_h1("6. Proposed Rollout")
add_table(
    headers=["Week", "Deliverable"],
    rows=[
        ["1", "Drizzle schema for catalog.initiatives + child tables; migrations applied to labs_ai_hub."],
        ["1–2", "apps/api/src/modules/roi/ scaffolded with Okta middleware and basic CRUD."],
        ["2", "Vitest suite on roi.calc; locked baseline numbers."],
        ["2", "Component ported into /roi-planner; window.storage replaced with API hooks."],
        ["2–3", "File decomposition; UI polish to match portal design system."],
        ["3+", "Pilot with one team (Growth Protocol or AI Engineering); gather feedback."],
        ["4+", "Open to org; add Project Catalogue cross-linking; vendor auto-population."],
    ],
    col_widths=[0.8, 5.9],
)

# ---------- 7. TUESDAY SHOWCASE ----------
add_h1("7. What I Will Showcase on Tuesday")
add_para("A 15-minute walkthrough in this order:")
add_bullet("1. Open the prototype with sample CoAction initiatives loaded (Vega, CoSave, Cortex, etc.).")
add_bullet("2. Executive view — show the headline FTE + BPO + Roadmap split. Click into 'Build vs Subscribe' tile to demonstrate filtering.")
add_bullet("3. Drill-down — click any initiative tile to show the calculation methodology panel. This is the trust-builder for finance.")
add_bullet("4. Run-rate view — switch year filter, point out Active / Planned / Roadmap pipeline.")
add_bullet("5. Investment arc — show the breakeven marker and cumulative net.")
add_bullet("6. HC harvest — show quarterly FTE positions, group by domain, expand business function.")
add_bullet("7. Roadmap Gantt — show savings phases highlighted in green.")
add_bullet("8. Goals tab — show planned vs realized progress against a sample goal.")
add_bullet("9. Add an initiative live to demonstrate the form and instant ROI preview.")
add_bullet("10. Flip to view-only mode to demonstrate the permission model.")
add_para("Then 10 minutes on integration plan, risks, and rollout. Ask for approval to proceed with Week 1 deliverables.")

# ---------- 8. ASKS ----------
add_h1("8. What I Need From This Meeting")
add_bullet("Approval to scaffold apps/api/src/modules/roi/ on the agreed Node + Express + Drizzle + Okta stack.")
add_bullet("Approval for the schema mapping into catalog / analytics / integration / lifecycle / engagement.")
add_bullet("A pilot team commitment (Growth Protocol or AI Engineering preferred — both already in Project Catalogue).")
add_bullet("Agreement that financial math gets Vitest coverage before any exec sees the numbers.")
add_bullet("Naming sign-off: 'ROI Planner' vs 'AI Investment Tracker' vs alternative.")

# ---------- FOOTER ----------
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
run = p.add_run("Prepared by Ramana Narayanan  |  CoAction Labs — AI Engineering")
run.font.size = Pt(9)
run.font.italic = True
run.font.color.rgb = GRAY

out = "/Users/navi/Work/smart-agent-foundry-main/docs/ROI_Planner_Tuesday_Brief.docx"
doc.save(out)
print("WROTE:", out)
